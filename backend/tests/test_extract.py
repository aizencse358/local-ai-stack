import io

import pytest
from docx import Document

import app.extract as extract_module
from app.extract import extract_text


def test_txt_and_md_decode_as_utf8():
    assert extract_text("notes.txt", "hello world".encode("utf-8")) == "hello world"
    assert extract_text("notes.md", "# heading".encode("utf-8")) == "# heading"


def test_pdf_joins_page_text(monkeypatch):
    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        def __init__(self, stream):
            self.pages = [FakePage("Page one"), FakePage("Page two")]

    monkeypatch.setattr(extract_module, "PdfReader", FakeReader)

    result = extract_text("doc.pdf", b"irrelevant bytes")

    assert result == "Page one\nPage two"


def test_pdf_reader_failure_raises_value_error(monkeypatch):
    def boom(stream):
        raise RuntimeError("corrupt PDF")

    monkeypatch.setattr(extract_module, "PdfReader", boom)

    with pytest.raises(ValueError):
        extract_text("broken.pdf", b"not a real pdf")


def test_docx_joins_paragraph_text():
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_text("doc.docx", buffer.getvalue())

    assert result == "First paragraph\nSecond paragraph"


def test_unsupported_extension_raises_value_error():
    with pytest.raises(ValueError):
        extract_text("file.xyz", b"whatever")
